from codecs import ignore_errors
from operator import ne
import pandas as pd
import json
import os
import re
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX

from pylatex import Document, Section, Subsection, Tabular, TextColor, Package, NoEscape
from pylatex import Math, TikZ, Axis, Plot, Figure, Matrix, Alignat
from pylatex.utils import italic
from pylatexenc import latexencode
import os


def main():
    """
    find findings in mimic (done), open-i, pubmed, and covid_chestxray
    """
    # file = f'metadata.csv'
    # pwd = os.getcwd()
    # filepath = os.path.join('/', *pwd.split(os.path.sep)[:-1], f'bert-AAD/data/nih_train.csv')
    file = f'indiana_reports.csv'
    # pubmed has no intersection with covid_chestxray
    # intersection between covid_chestxray and open-i is ['tuberculosis', 'pneumonia'] + normal/no finding
    covid_findings = ['Aspergillosis', 'Aspiration', 'Bacterial', 'COVID-19',
       'Chlamydophila', 'E.Coli', 'Fungal', 'H1N1', 'Herpes ',
       'Influenza', 'Klebsiella', 'Legionella', 'Lipoid', 'MERS-CoV',
       'MRSA', 'Mycoplasma', 'No Finding', 'Nocardia', 'Pneumocystis',
       'Pneumonia', 'SARS', 'Staphylococcus', 'Streptococcus',
       'Tuberculosis', 'Unknown', 'Varicella', 'Viral', 'todo']
    covid_findings_split = []
    haha = [covid_findings_split.extend(item.split("/")) for item in covid_findings]
    covid_findings_unique = np.unique(covid_findings_split)
    covid_findings_unique = [item.lower() for item in covid_findings_unique]
    df = pd.read_csv(file)  # [930, 30]
    findings = df['Problems']
    reports = df['findings']  # if nan use df['impression']
    all_findings = pd.unique(findings)
    unique_findings = []
    hh = [unique_findings.extend(item.split(';')) for item in all_findings]
    unique_findings = np.unique(unique_findings)
    unique_findings = [item.lower() for item in unique_findings]
    intersect = list(set(covid_findings_unique).intersection(set(unique_findings)))
    print(intersect)


def filter_covid():
    """
    intersection between covid_chestxray and open-i is ['tuberculosis', 'pneumonia'] + normal/no finding
    In covid_chestxray, find report + label in above findings ['No Finding']
    In open-i, find report + label in above findings ['normal']
    done
    """
    covid = f'metadata.csv'
    covid_df = pd.read_csv(covid)
    findings = covid_df['finding']
    reports = covid_df['clinical_notes']
    columns=['tuberculosis', 'pneumonia', 'no finding', 'report']
    covid_new_df = pd.DataFrame(columns=columns)
    for i, row in covid_df.iterrows():
        fins = row['finding'].lower().split('/')  # findings
        new_row = [1 if item in fins else 0 for item in columns[:-1]]
        if sum(new_row) > 0:
            new_row.append(row['clinical_notes'])
            new_row = pd.Series(new_row, index=covid_new_df.columns)
            covid_new_df = covid_new_df.append(new_row, ignore_index=True)
    # remove duplicated rows
    covid_new_df = covid_new_df.drop_duplicates(subset=['report'])
    covid_new_df.dropna(subset=['report'], inplace=True)  # [612, 4]
    covid_new_df.to_csv(f'data/covid_all.csv')
    covid_train_df = covid_new_df.sample(frac=0.81, replace=False)
    covid_train_df.to_csv(f'data/covid_train.csv')


def filter_openi():
    """
    intersection between covid_chestxray and open-i is ['tuberculosis', 'pneumonia'] + normal/no finding
    In covid_chestxray, find report + label in above findings ['No Finding']
    In open-i, find report + label in above findings ['normal']
    """
    openi = f'data/indiana_reports.csv'
    df = pd.read_csv(openi)
    findings = df['Problems']
    unique_findings = []
    haha = [unique_findings.extend(item.split(';')) for item in findings]
    unique_findings = list(set(unique_findings))
    print(unique_findings)
    reports = df['findings']
    df['report'] = (df['findings'] + df['impression']).astype(str)
    columns=['tuberculosis', 'pneumonia', 'normal', 'report']
    openi_new_df = pd.DataFrame(columns=columns)
    for i, row in df.iterrows():
        fins = row['Problems'].lower().split(';')
        new_row = [1 if item in fins else 0 for item in columns[:-1]]
        if sum(new_row) > 0:
            new_row.append(row['report'])
            new_row = pd.Series(new_row, index=openi_new_df.columns)
            openi_new_df = openi_new_df.append(new_row, ignore_index=True)

    openi_new_df = openi_new_df.drop_duplicates(subset=['report'])
    openi_new_df.dropna(subset=['report'], inplace=True)
    openi_new_df.to_csv(f'data/openi_all.csv')
    openi_train_df = openi_new_df.sample(frac=0.8, replace=False)
    openi_train_df.to_csv(f'data/openi_train.csv')


def read_json():
    """
    failed to add paragraph into row_cell, don't use
    """
    # read from findings.txt
    with open(f'data/findings.txt', 'r') as f1:
        findings_list = f1.read().split('\n')
        # print(findings_list)  # correct

    file = f'data/06102020.litcovid.released.text.json'
    with open(file) as f:
        data = json.load(f)
        reports = []
        docu = Document()
        docu.add_heading('COVID CXR CT reports', 0)
        table = docu.add_table(rows=1, cols=3, style='TableGrid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Id'
        hdr_cells[1].text = 'Report'
        hdr_cells[2].text = 'Label'
        for i, item in enumerate(data):
            if 'text' in item:
                report = ""
                for infons in item['text']:
                    report += infons['text']
                reports.append(report)

        reports = set(reports)
        for i, rep in enumerate(reports):
            z1 = re.findall(r"(?=("+'|'.join(findings_list)+r"))", rep, flags=re.I)
            result = re.finditer(r"(?=("+'|'.join(findings_list)+r"))", rep, flags=re.I)
            
            if len(z1) > 1:
                indices = []
                for match in result:
                    indices.append(match.span(1))  # span(1): (1965, 1974)
                    # print(rep[match.span(1)[0]:match.span(1)[1]])
                para = docu.add_paragraph(rep[:indices[0][0]])
                for j, (idx1, idx2) in enumerate(indices):
                    para.add_run(rep[idx1:idx2]).font.highlight_color = WD_COLOR_INDEX.YELLOW
                    if j+1 < len(indices):
                        para.add_run(rep[idx2:indices[j+1][0]])
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].add_paragraph(style=None)
                row_cells[1].paragraphs[0] = para
                row_cells[2].text = str({word:1 for word in z1})

        docu.save('data/COVID_CXR_CT.docx')


def latex_example():
    """
    TODO: convert above function to generate latex
    """
    # image_filename = os.path.join(os.path.dirname(__file__), 'kitten.jpg')
  
    geometry_options = {"tmargin": "1cm", "lmargin": "1cm"}
    doc = Document(geometry_options=geometry_options)
  
    # creating a pdf with title "the simple stuff"
    # with doc.create(Section('The simple stuff')):
    #     doc.append('Some regular text and some')
    #     doc.append(italic('italic text. '))
    #     doc.append('\nAlso some crazy characters: $&#{}')
    #     with doc.create(Subsection('Math that is incorrect')):
    #         doc.append(Math(data=['2*3', '=', 9]))
  
    #     # creating subsection of a pdf
    #     with doc.create(Subsection('Table of something')):
    #         with doc.create(Tabular('rc|cl')) as table:
    #             table.add_hline()
    #             table.add_row((1, 2, 3, 4))
    #             table.add_hline(1, 2)
    #             table.add_empty_row()
    #             table.add_row((4, 5, 6, 7))
  
    #  # making a pdf using .generate_pdf
    # doc.generate_pdf(f'data/full', clean_tex=False)
    # tex = doc.dumps()  # The document as string in LaTeX syntax
    # doc = Document()
    doc.packages.append(Package('color'))
    # doc.preamble.append(NoEscape(r'\newcommand{\exampleCommand}[3]{\color{#1} #2 #3 \color{black}}'))
    with doc.create(Tabular('c|c|c')) as table:
        table.add_hline()
        hl_report = NoEscape(r'haha \textcolor{red}{Hello World!}  xixi ')
        hl_report += NoEscape(r'xiha')
        table.add_row((4, hl_report, 6))
        table.add_hline()

    # doc.preamble.append(NoEscape(r'\newcommand*\Entry[2]{\sffamily#1 & #2}'))
    # doc.append(NoEscape(r'\begin{tabular}{ll}'))
    # doc.append(NoEscape(r'\Entry{abc}{123}\\'))
    # doc.append(NoEscape(r'\end{tabular}'))

    doc.generate_pdf(f'data/pycommands', clean_tex=False)


def json_latex():
    """
    TODO: convert above function to generate latex
    """
    geometry_options = {"tmargin": "1cm", "lmargin": "1cm"}
    doc = Document(geometry_options=geometry_options)
    with open(f'data/findings.txt', 'r') as f1:
        findings_list = f1.read().split('\n')
        # print(findings_list)  # correct

    file = f'data/06102020.litcovid.released.text.json'
    with open(file) as f:
        data = json.load(f)
        reports = []
        for i, item in enumerate(data):
            if 'text' in item:
                report = ""
                for infons in item['text']:
                    report += infons['text']
                reports.append(report)

        reports = sorted(list(set(reports)))
        doc.packages.append(Package('color'))
        
        for i, rep in enumerate(reports):
            with doc.create(Section(f'{i}')):
                z1 = re.findall(r"(?=("+'|'.join(findings_list)+r"))", rep, flags=re.I)
                result = re.finditer(r"(?=("+'|'.join(findings_list)+r"))", rep, flags=re.I)
                new_report = ""
                if len(z1) > 1:
                    indices = []  # keywords indices
                    for match in result:
                        indices.append(match.span(1))  # span(1): (1965, 1974)

                    new_report += rep[:indices[0][0]]
                    for j, (idx1, idx2) in enumerate(indices):
                        new_report += '\\textcolor{red}{' + rep[idx1:idx2] + '}'
                        if j+1 < len(indices):
                            new_report += rep[idx2:indices[j+1][0]]
                    new_report += rep[idx2:]
                # build highlighed report in hl_report
                # new_report = new_report.encode('utf-8','ignore').decode("utf-8")
                doc.append(NoEscape(new_report))

                with doc.create(Subsection('Labels')):
                    doc.append(NoEscape(str({word:1 for word in z1})))

    # doc.generate_pdf(f'data/COVID_CXR_CT', clean_tex=False)
    doc.generate_tex(f'data/COVID_CXR_CT')
    tex = doc.dumps()


def read_xinyue_label():
    """
    1. use labels from xinyue's result, 
    2. randomly sample 100, make tex that highlight keywords in the report
    """
    label_file = f'data/covid_labels.json'
    with open(label_file) as f2:
        labels = json.load(f2)
        # print(labels)

    file = f'data/06102020.litcovid.released.text.json'
    with open(file) as f1:
        data = json.load(f1)
        reports = []
        for i, item in enumerate(data):
            if 'text' in item:
                report = ""
                for infons in item['text']:
                    report += infons['text']
                reports.append(report)

        reports = sorted(list(set(reports)))
        geometry_options = {"tmargin": "1cm", "lmargin": "1cm"}
        doc = Document(geometry_options=geometry_options)
        doc.packages.append(Package('color'))
        intersected_findings = []
        class_counter = {}
        for i, (rep, label) in enumerate(zip(reports, labels)):
            findings_list = list(label['entity'].keys())
            with doc.create(Section(f'{i+1}')):
                if len(findings_list) > 0:
                    for fi in findings_list:
                        if fi in class_counter:
                            class_counter[fi] += 1
                        else:
                            class_counter[fi] = 1
                    intersected_findings.extend(findings_list)
                    z1 = re.findall(r"(?=("+'|'.join(findings_list)+r"))", rep, flags=re.I)
                    result = re.finditer(r"(?=("+'|'.join(findings_list)+r"))", rep, flags=re.I)
                    new_report = ""
                    if len(z1) > 1:
                        indices = []  # keywords indices
                        for match in result:
                            indices.append(match.span(1))  # span(1): (1965, 1974)

                        new_report += rep[:indices[0][0]]
                        for j, (idx1, idx2) in enumerate(indices):
                            new_report += '\\textcolor{red}{' + rep[idx1:idx2] + '}'
                            if j+1 < len(indices):
                                new_report += rep[idx2:indices[j+1][0]]
                        new_report += rep[idx2:]

                        doc.append(NoEscape(new_report))

                        with doc.create(Subsection('Labels')):
                            doc.append(NoEscape(str({word:1 for word in findings_list})))
                else:
                    # no finding
                    doc.append(NoEscape(rep))
                    doc.create(Subsection('Labels'))

        print(set(intersected_findings))
        print(class_counter)
        doc.generate_tex(f'data/COVID_CXR_CT')
        tex = doc.dumps()


def build_train_all_csv():
    """
    from report and label build df with report + multi-hot label
    """
    unique_labels = ['consolidation', 'edema', 'cardiomegaly', 'pleural effusion', 
    'Lung Lesion', 'fracture', 'pneumothorax', 'pneumonia', 'medical device', 
    'lung opacity', 'atelectasis']
    label_file = f'data/covid_labels.json'
    with open(label_file) as f2:
        labels = json.load(f2)


    file = f'data/06102020.litcovid.released.text.json'
    with open(file) as f1:
        data = json.load(f1)
        reports = []
        for i, item in enumerate(data):
            if 'text' in item:
                report = ""
                for infons in item['text']:
                    report += infons['text']
                reports.append(report)
        df_cols = unique_labels + ['report']
        df = pd.DataFrame(columns=df_cols)
        reports = sorted(list(set(reports)))
        for i, (rep, label) in enumerate(zip(reports, labels)):            
            empty_labels = np.zeros(len(unique_labels))
            findings_list = list(label['entity'].keys())
            for fi in findings_list:
                idx = unique_labels.index(fi)
                empty_labels[idx] = 1

            row = empty_labels.tolist() + [rep]
            row = pd.Series(row, index=df.columns)
            df = df.append(row, ignore_index=True)
    train_df = df.sample(frac=0.8, replace=False)
    train_df.to_csv(f'output/ctcxr_train.csv')
    df.to_csv(f'output/ctcxr_all.csv')


if __name__ == '__main__':
    # main()
    # filter_covid()
    # filter_openi()
    # read_json()
    # latex_example()
    # json_latex()
    # read_xinyue_label()
    build_train_all_csv()
